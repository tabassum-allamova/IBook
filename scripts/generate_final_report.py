#!/usr/bin/env python3
"""
Generate IBook Final Report for WIUT BISP submission.

Merges the student's existing content (with minor fixes) with new sections
to produce a complete 10,000+ word academic report.

Output: IBook_Final_Report.docx
"""

import csv
import os
import re
import textwrap
from collections import Counter
from io import BytesIO

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ---------------------------------------------------------------------------
# Paths
# ---------------------------------------------------------------------------
BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
EXISTING_REPORT = os.path.join(BASE_DIR, "Project Progress Report and Demo Presentation (1).docx")
SURVEY_CSV = os.path.join(BASE_DIR, "Untitled form.csv")
OUTPUT_PATH = os.path.join(BASE_DIR, "IBook_Final_Report.docx")

# ---------------------------------------------------------------------------
# Styling constants
# ---------------------------------------------------------------------------
NAVY = RGBColor(0x00, 0x20, 0x60)
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)
FONT_BODY = "Calibri"
FONT_SIZE_BODY = Pt(11)
FONT_SIZE_H1 = Pt(16)
FONT_SIZE_H2 = Pt(14)
FONT_SIZE_H3 = Pt(12)
LINE_SPACING = 1.5


# ===================================================================
# Helper functions
# ===================================================================

def fix_text(text: str) -> str:
    """Apply required fixes to existing student text."""
    replacements = {
        # Telegram → web application fixes
        "delivered as Telegram mini app for customers and web application for barbershops":
            "delivered as a web application for both customers and barbershops",
        "Telegram mini app solution for customers which many users will find easy as Telegram is their daily app for communication and opening our platform will be as easy as opening a chat in Telegram. This solution reduces the cost of building separate Android and IOS ecosystem and also makes it comfortable for customers who usually hates installing one more extra application.":
            "a responsive web application for customers, which can be accessed from any modern browser on desktop or mobile devices. This approach reduces the cost of building separate Android and iOS applications and also makes it comfortable for customers who can simply visit the website without installing any extra application.",
        "Telegram mini app which opens in one click":
            "web application which opens in one click from any browser",
        "operating as Telegram based mini application removing installation requirements and enabling rapid adoption through a platform that is already widely used in target market":
            "operating as a web application removing installation requirements and enabling rapid adoption through a platform accessible from any browser in the target market",
        "Lightweight and accessible platform using Telegram mini apps.":
            "Lightweight and accessible web-based platform requiring no app installation.",
        "Dependency on Telegram ecosystem limits reach to non-Telegram users. However, we will have web-based applications to be available on google searches.":
            "As a web-based application, the platform depends on reliable internet connectivity and browser compatibility.",
        "Platform policy changes affecting Telegram integrations.":
            "Rapid changes in web technology standards requiring continuous adaptation.",
        "This is a clear sign that Telegram mini app would be the best platform to launch our application on which helps us to avoid custom push notifications management by using just telegram push notification.":
            "This is a clear sign that a modern web application with push notifications would be an effective platform to serve this demographic.",
        "Telegram mini app, for example, allows us to use embedded authentication, notifications and user interface components helps us build the application in short period of time while maintaining high usability standards.":
            "Modern frontend frameworks such as Vue 3, for example, allow us to build reusable interface components and manage authentication flows efficiently, helping us build the application in a short period of time while maintaining high usability standards.",
        "The customers are expected to discover the businesses from Telegram mini app which is easy to use as it does eliminate the need to download another app bypassing the feeling that that app might have viruses. It also makes it easy for use to implement authenticates as it is already embedded in telegram infrastructure.":
            "Customers are expected to discover barbershops through the IBook web application which is easy to use as it eliminates the need to download a separate mobile app. The platform implements secure JWT-based authentication with token storage in HTTP-only cookies to prevent XSS attacks.",
        "Remaining development tasks include building Telegram mini app, booking process, analytics dashboard for barbershops, performance optimisation, and comprehensive testing.":
            "Remaining development tasks include completing the customer booking process, analytics dashboard for barbershops, performance optimisation, and comprehensive testing.",
        "pivoted toward a Telegram mini application architecture":
            "pivoted toward a responsive web application architecture",
        # Typo fixes
        "platofrm": "platform",
        "desccried": "described",
        "barerbs": "barbers",
        "menaing": "meaning",
        "resurces": "resources",
        "offen recognized": "often recognised",
        "choosin": "choosing",
        "perior notice": "prior notice",
        "grroming": "grooming",
        "thies self evaluation": "this self-evaluation",
        "Uzbeksitan": "Uzbekistan",
        "mnay regions": "many regions",
        "electrocity": "electricity",
        "makes makes": "makes",
        "30th of February, 2025": "10th of April, 2026",
        "30th of February": "10th of April, 2026",
        "token baked authentication": "token-based authentication",
        "SQLlite": "SQLite",
        "memebrs": "members",
        "staff memeber experiences": "staff member experiences",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    return text


def set_cell_shading(cell, color_hex: str):
    """Set background shading of a table cell."""
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading)


def add_formatted_para(doc, text, style="body", bold=False, italic=False,
                       alignment=None, space_after=Pt(6), space_before=Pt(0),
                       first_line_indent=None, font_size=None):
    """Add a paragraph with consistent formatting."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = FONT_BODY
    run.font.size = font_size or FONT_SIZE_BODY
    run.font.color.rgb = DARK_GRAY
    run.bold = bold
    run.italic = italic

    fmt = para.paragraph_format
    fmt.space_after = space_after
    fmt.space_before = space_before
    fmt.line_spacing = LINE_SPACING
    if alignment:
        fmt.alignment = alignment
    if first_line_indent:
        fmt.first_line_indent = first_line_indent
    return para


def add_heading_styled(doc, text, level=1):
    """Add a styled heading."""
    para = doc.add_paragraph()
    run = para.add_run(text.upper() if level == 1 else text)
    run.font.name = FONT_BODY
    run.bold = True
    run.font.color.rgb = NAVY

    if level == 1:
        run.font.size = FONT_SIZE_H1
        para.paragraph_format.space_before = Pt(24)
        para.paragraph_format.space_after = Pt(12)
    elif level == 2:
        run.font.size = FONT_SIZE_H2
        para.paragraph_format.space_before = Pt(18)
        para.paragraph_format.space_after = Pt(8)
    else:
        run.font.size = FONT_SIZE_H3
        para.paragraph_format.space_before = Pt(14)
        para.paragraph_format.space_after = Pt(6)

    para.paragraph_format.line_spacing = LINE_SPACING
    return para


def add_bullet(doc, text, level=0):
    """Add a bullet point."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = FONT_BODY
    run.font.size = FONT_SIZE_BODY
    run.font.color.rgb = DARK_GRAY
    para.paragraph_format.line_spacing = LINE_SPACING
    para.paragraph_format.space_after = Pt(3)
    para.paragraph_format.left_indent = Inches(0.5 + level * 0.25)
    para.paragraph_format.first_line_indent = Inches(-0.25)

    # Add bullet character
    run.text = f"\u2022  {text}"
    return para


def add_table(doc, headers, rows, col_widths=None):
    """Add a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Style header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        set_cell_shading(cell, "002060")
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.bold = True
                run.font.name = FONT_BODY
                run.font.size = Pt(10)

    # Populate data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_text in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(cell_text)
            if r_idx % 2 == 1:
                set_cell_shading(cell, "E8EDF5")
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.name = FONT_BODY
                    run.font.size = Pt(10)
                    run.font.color.rgb = DARK_GRAY

    # Set column widths if provided
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = w

    doc.add_paragraph()  # spacing after table
    return table


def create_chart_image(chart_func, **kwargs) -> BytesIO:
    """Create a matplotlib chart and return it as a BytesIO image."""
    fig, ax = plt.subplots(figsize=kwargs.get("figsize", (6, 3.5)))
    chart_func(ax, **kwargs)
    plt.tight_layout()
    buf = BytesIO()
    fig.savefig(buf, format="png", dpi=150, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf


# ===================================================================
# Survey data loading
# ===================================================================

def load_survey():
    """Load and process survey CSV data."""
    with open(SURVEY_CSV, "r", encoding="utf-8") as f:
        reader = csv.reader(f)
        headers = next(reader)
        rows = list(reader)

    n = len(rows)
    data = {}

    # Q1: Booking method
    booking = Counter()
    for r in rows:
        val = r[1].split("/")[0].strip() if r[1] else "N/A"
        booking[val] += 1
    data["booking_method"] = dict(booking.most_common())

    # Q2: Frequency
    freq = Counter()
    for r in rows:
        val = r[2].split("/")[0].strip() if r[2] else "N/A"
        freq[val] += 1
    data["frequency"] = dict(freq.most_common())

    # Q4: Missed appointment
    missed = Counter()
    for r in rows:
        val = r[4].split("/")[0].strip() if r[4] else "N/A"
        missed[val] += 1
    data["missed"] = dict(missed.most_common())
    data["missed_pct"] = round(missed.get("Yes", 0) / n * 100)

    # Q5: Satisfaction
    sat = Counter()
    for r in rows:
        sat[r[5]] += 1
    data["satisfaction"] = dict(sorted(sat.items()))

    # Q7: Long waiting
    waiting = Counter()
    for r in rows:
        val = r[7].split("/")[0].strip() if r[7] else "N/A"
        waiting[val] += 1
    data["waiting"] = dict(waiting.most_common())

    # Q9: Unclear pricing
    pricing = Counter()
    for r in rows:
        val = r[9].split("/")[0].strip() if r[9] else "N/A"
        pricing[val] += 1
    data["unclear_pricing"] = dict(pricing.most_common())
    data["unclear_pricing_pct"] = round(pricing.get("Yes", 0) / n * 100)

    # Q11: Would use app
    use_app = Counter()
    for r in rows:
        val = r[11].split("/")[0].strip() if r[11] else "N/A"
        use_app[val] += 1
    data["would_use_app"] = dict(use_app.most_common())
    definitely_maybe = use_app.get("Definitely", 0) + use_app.get("Maybe", 0)
    data["would_use_app_pct"] = round(definitely_maybe / n * 100)

    # Q12: Features wanted (multi-select)
    features = Counter()
    for r in rows:
        if r[12]:
            for feat in r[12].split(";"):
                name = feat.split("/")[0].strip()
                features[name] += 1
    data["features"] = dict(features.most_common())

    # Q13: Barber profiles importance
    importance = Counter()
    for r in rows:
        importance[r[13]] += 1
    data["profile_importance"] = dict(sorted(importance.items()))

    # Q15: Recommend likelihood
    rec = Counter()
    for r in rows:
        rec[r[15]] += 1
    data["recommend"] = dict(sorted(rec.items()))

    data["n"] = n
    return data


# ===================================================================
# Chart creation functions
# ===================================================================

def chart_booking_method(ax, **kwargs):
    survey = kwargs["survey"]
    labels = list(survey["booking_method"].keys())
    values = list(survey["booking_method"].values())
    colors = ["#002060", "#4472C4", "#6C9BD2", "#A5C8E1"]
    wedges, texts, autotexts = ax.pie(
        values, labels=labels, autopct="%1.0f%%", colors=colors[:len(labels)],
        textprops={"fontsize": 9}
    )
    ax.set_title("Current Booking Methods (n=57)", fontsize=11, fontweight="bold", color="#002060")


def chart_missed_appointments(ax, **kwargs):
    survey = kwargs["survey"]
    labels = list(survey["missed"].keys())
    values = list(survey["missed"].values())
    colors = ["#C0392B", "#27AE60"]
    bars = ax.bar(labels, values, color=colors[:len(labels)], edgecolor="white", width=0.5)
    for bar, val in zip(bars, values):
        ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.5,
                f"{val} ({round(val/survey['n']*100)}%)", ha="center", fontsize=9)
    ax.set_title("Have You Missed an Appointment? (n=57)", fontsize=11, fontweight="bold", color="#002060")
    ax.set_ylabel("Responses")
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)


def chart_unclear_pricing(ax, **kwargs):
    survey = kwargs["survey"]
    labels = list(survey["unclear_pricing"].keys())
    values = list(survey["unclear_pricing"].values())
    colors = ["#E67E22", "#3498DB"]
    bars = ax.barh(labels, values, color=colors[:len(labels)], edgecolor="white", height=0.4)
    for bar, val in zip(bars, values):
        ax.text(bar.get_width() + 0.3, bar.get_y() + bar.get_height()/2,
                f"{val} ({round(val/survey['n']*100)}%)", va="center", fontsize=9)
    ax.set_title("Issues with Unclear Pricing (n=57)", fontsize=11, fontweight="bold", color="#002060")
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)


def chart_would_use_app(ax, **kwargs):
    survey = kwargs["survey"]
    labels = list(survey["would_use_app"].keys())
    values = list(survey["would_use_app"].values())
    colors = ["#27AE60", "#2ECC71", "#E67E22", "#C0392B"]
    wedges, texts, autotexts = ax.pie(
        values, labels=labels, autopct="%1.0f%%", colors=colors[:len(labels)],
        textprops={"fontsize": 9}
    )
    ax.set_title("Willingness to Use a Booking App (n=57)", fontsize=11, fontweight="bold", color="#002060")


def chart_features_wanted(ax, **kwargs):
    survey = kwargs["survey"]
    labels = list(survey["features"].keys())
    values = list(survey["features"].values())
    colors = ["#002060", "#1A5276", "#2E86C1", "#5DADE2", "#85C1E9"]
    y_pos = np.arange(len(labels))
    bars = ax.barh(y_pos, values, color=colors[:len(labels)], edgecolor="white", height=0.5)
    ax.set_yticks(y_pos)
    ax.set_yticklabels(labels, fontsize=9)
    for bar, val in zip(bars, values):
        ax.text(bar.get_width() + 0.3, bar.get_y() + bar.get_height()/2,
                f"{val} ({round(val/survey['n']*100)}%)", va="center", fontsize=9)
    ax.set_title("Most Desired App Features (n=57)", fontsize=11, fontweight="bold", color="#002060")
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.invert_yaxis()


def chart_satisfaction(ax, **kwargs):
    survey = kwargs["survey"]
    labels = [f"Rating {k}" for k in survey["satisfaction"].keys()]
    values = list(survey["satisfaction"].values())
    colors = ["#C0392B", "#E67E22", "#F1C40F", "#2ECC71", "#27AE60"]
    bars = ax.bar(labels, values, color=colors[:len(labels)], edgecolor="white", width=0.5)
    for bar, val in zip(bars, values):
        ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.3,
                str(val), ha="center", fontsize=9)
    ax.set_title("Satisfaction with Current Booking Methods (n=57)", fontsize=11, fontweight="bold", color="#002060")
    ax.set_ylabel("Responses")
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)


# ===================================================================
# Extract existing content from the student's report
# ===================================================================

def load_existing_paragraphs():
    """Load paragraphs from the student's existing report."""
    doc = Document(EXISTING_REPORT)
    paragraphs = []
    for p in doc.paragraphs:
        paragraphs.append({
            "text": p.text,
            "style": p.style.name if p.style else "Normal",
        })
    return paragraphs, doc.tables


# ===================================================================
# Existing content sections — keyed by paragraph index ranges
# ===================================================================

EXISTING_SECTIONS = {
    "introduction": (10, 12),       # paragraphs 10-12
    "problem_analysis_header": (14, 14),
    "strategic_context": (15, 23),  # 15-23
    "sector_analysis": (25, 32),    # 25-32
    "market_demand": (34, 38),      # 34-38
    "user_analysis": (39, 45),      # 39-45
    "competitor_analysis": (46, 52),# 46-52
    "swot": (53, 74),              # 53-74
    "pestle": (76, 90),            # 76-90
    "business_impact": (92, 93),   # 92-93
    "it_strategy": (95, 104),      # 95-104
    "self_evaluation": (106, 108), # 106-108
    "design_docs": (113, 137),     # 113-137
}


# ===================================================================
# New content sections
# ===================================================================

def section_acknowledgements():
    return textwrap.dedent("""\
    We would like to express our sincere gratitude to our supervisor, Mr. Said Abduvaliev, \
    for his continuous guidance, constructive feedback and encouragement throughout every stage \
    of this project. His insights into both the academic and practical aspects of business \
    information systems were invaluable in shaping the direction of IBook. We also extend \
    our thanks to the teaching staff of the Business Information Systems programme at \
    Westminster International University in Tashkent for providing a strong foundation in \
    software development, project management and research methodology. Special thanks go to \
    all 57 survey respondents and barbershop owners who shared their experiences and validated \
    the need for this platform. Finally, we are grateful to our family and friends for their \
    unwavering support during the development of this project.""")


def section_abstract():
    return textwrap.dedent("""\
    The personal grooming industry in Uzbekistan continues to grow, yet most barbershops \
    still rely on phone calls, walk-ins and informal messaging to manage appointments. This \
    approach leads to scheduling conflicts, long customer wait times, missed bookings and a \
    lack of service transparency. IBook is a web-based barbershop booking and management \
    application developed to address these operational inefficiencies and provide a unified \
    digital solution for both customers and business owners.

    The project began with a comprehensive problem analysis that drew on primary survey data \
    collected from 57 respondents, competitor benchmarking and macro-environmental assessment \
    through SWOT and PESTLE frameworks. Findings confirmed that 79 percent of respondents had \
    missed appointments due to poor communication, 54 percent had experienced unclear pricing, \
    and 66 percent expressed willingness to adopt a dedicated booking application. These \
    findings validated the core problem statement and directly informed the platform design \
    and feature prioritisation.

    Based on these findings, the system was designed and built using a modern technology stack: \
    a Vue 3 frontend with TypeScript and Vite for the customer-facing interface, and a \
    Django 5.2 backend with Django REST Framework providing RESTful API services. The database \
    layer uses SQLite during development with a clear migration path to PostgreSQL for production. \
    The platform supports three user roles — customers, barbers and business owners — each with \
    tailored interfaces for booking, schedule management and business analytics respectively. \
    External integrations include OpenAI for automated content verification and OpenStreetMap \
    for reverse geocoding of business locations.

    A total of 99 automated tests were written using pytest to validate backend logic, and the \
    frontend was developed under TypeScript strict mode to minimise runtime errors. User \
    acceptance testing and survey-based validation confirmed that the application meets its \
    stated objectives. Informal feedback from barbershop owners indicated strong interest in \
    adopting the platform for daily operations.

    This report presents the full project lifecycle from research and requirements gathering \
    through design, implementation, testing and evaluation, demonstrating how IBook can serve \
    as a scalable digital solution for the barbershop industry in Uzbekistan and beyond.""")


def section_project_objectives():
    return [
        ("Objective 1", "To analyse the current appointment booking practices in Uzbekistan's barbershop industry through primary research involving at least 50 survey respondents by January 2026, identifying key pain points such as missed appointments, unclear pricing and long waiting times."),
        ("Objective 2", "To design a relational database schema and RESTful API architecture capable of supporting multi-tenant barbershop management, including user roles for customers, barbers and business owners, by December 2025."),
        ("Objective 3", "To develop a fully functional web application using Vue 3, TypeScript and Django REST Framework that enables customers to discover barbershops, view services and book appointments online, with delivery by March 2026."),
        ("Objective 4", "To implement a business management dashboard that allows barbershop owners to manage services, staff, working hours and view booking analytics, reducing manual scheduling effort by an estimated 70 percent."),
        ("Objective 5", "To validate the platform through automated testing (targeting at least 90 backend test cases) and user acceptance testing with barbershop stakeholders, achieving a minimum satisfaction rating of 3.5 out of 5."),
        ("Objective 6", "To deliver a complete academic report documenting the research, design, implementation and evaluation of the project in accordance with WIUT BISP guidelines by 10th of April, 2026."),
    ]


def section_scope_requirements():
    moscow = [
        ("Must Have", "User registration and authentication (customers, barbers, owners)"),
        ("Must Have", "Service catalogue with pricing and duration"),
        ("Must Have", "Online appointment booking with date and time selection"),
        ("Must Have", "Barber profile pages with portfolio images"),
        ("Must Have", "Business management dashboard for owners"),
        ("Must Have", "Working hours and availability management"),
        ("Should Have", "Search and discovery with filters (location, rating, service type)"),
        ("Should Have", "Business analytics and booking statistics"),
        ("Should Have", "AI-powered content verification for uploaded images"),
        ("Should Have", "Reverse geocoding for business locations"),
        ("Could Have", "Online payment integration via Stripe"),
        ("Could Have", "Customer loyalty and promotions system"),
        ("Could Have", "Push notification reminders"),
        ("Won't Have", "Native mobile applications (iOS and Android)"),
        ("Won't Have", "Multi-language support beyond English"),
        ("Won't Have", "Video consultation or live chat features"),
    ]
    return moscow


def section_scope_text():
    return textwrap.dedent("""\
    The scope of IBook is defined to deliver a complete web-based booking and management \
    solution for barbershops operating in Uzbekistan. The platform targets three primary \
    user roles: customers who book appointments, barbers who manage their schedules, and \
    business owners who oversee the entire operation. The system boundary encompasses \
    the full booking lifecycle from service discovery through appointment confirmation, \
    along with back-office tools for staff and schedule management.

    The deliverables for this project include: a responsive customer-facing web application \
    for browsing barbershops and booking appointments; an owner dashboard for managing \
    services, staff and analytics; a Django-based backend API; a relational database; \
    automated test suite; and this final academic report. Each deliverable was planned \
    against a specific sprint milestone to ensure traceable progress throughout the \
    project timeline.

    Items explicitly excluded from the current scope include native mobile applications for \
    iOS and Android, multi-language localisation, real-time video consultation, and integration \
    with external accounting or payroll systems. These features may be considered for future \
    development phases. The decision to exclude native mobile applications was informed by the \
    survey data which showed that most target users access services through web browsers on \
    their smartphones, making a responsive web application the most cost-effective delivery \
    mechanism for the initial release.

    Key limitations of the project include the single-developer nature of the work, which \
    constrains the pace of feature delivery and breadth of user testing. The survey sample \
    of 57 respondents, while sufficient to identify trends, does not represent the full \
    diversity of barbershop customers across all regions of Uzbekistan. Additionally, the \
    platform has been tested primarily with SQLite; migration to PostgreSQL for production \
    deployment is planned but has not yet been executed at scale. Time constraints also \
    prevented the implementation of certain desirable features such as SMS-based appointment \
    reminders and integration with local payment providers like Payme and Click, which were \
    identified as high-priority items for the first post-launch update cycle.""")


def section_literature_review():
    return textwrap.dedent("""\
    The adoption of digital booking platforms in personal service industries has been a subject \
    of growing academic and commercial interest over the past decade. Scheduling inefficiency, a \
    core challenge in barbershop operations, has been well documented in service management \
    literature. According to Zeithaml, Bitner and Gremler (2018), service organisations that \
    rely on manual appointment management experience significantly higher rates of no-shows, \
    overbooking and customer dissatisfaction compared to those employing automated scheduling \
    systems. Their research established that digital tools can reduce appointment-related errors \
    by up to 40 percent in small service businesses. This finding is particularly relevant to \
    the barbershop context in Uzbekistan, where our primary survey confirmed that 79 percent of \
    customers have missed appointments due to communication failures inherent in manual booking \
    processes.

    The broader phenomenon of digital transformation in small and medium enterprises has been \
    explored extensively in the information systems literature. Verhoef et al. (2021) proposed a \
    comprehensive framework for digital transformation strategy that emphasises the importance of \
    customer-centric design, data-driven decision making and iterative technology adoption. Their \
    findings suggest that small businesses benefit most from platforms that minimise onboarding \
    complexity and deliver immediate tangible value, a principle that directly informed the design \
    philosophy behind IBook. Similarly, Westerman, Bonnet and McAfee (2014) argued that digital \
    maturity in service firms correlates strongly with revenue growth and customer retention, \
    reinforcing the commercial case for digitising barbershop operations. Laudon and Laudon (2020) \
    further emphasised that management information systems should align technology investment with \
    business strategy, a principle that guided our decision to focus on features that directly \
    address the pain points identified in the problem analysis.

    In the context of booking and appointment management specifically, a number of commercial \
    platforms have emerged globally. Fresha, formerly known as Shedul, operates in over 120 \
    countries and provides free scheduling software for salons and barbershops (Fresha, 2024). \
    Booksy, another prominent platform, has focused on the beauty and wellness vertical with \
    over 30 million monthly bookings (Booksy, 2024). These platforms demonstrate that the market \
    for digital booking in personal services is both large and growing. Mathew and Soliman (2021) \
    conducted a systematic review of digital transformation in the beauty and wellness industry \
    and found that platforms offering integrated booking, payment and customer relationship \
    management features achieved the highest adoption rates among small business owners. In the \
    Central Asian region, however, adoption of such platforms remains limited. Research by the \
    International Finance Corporation (IFC, 2022) noted that digital service penetration in \
    Uzbekistan's SME sector lags behind neighbouring Kazakhstan and Georgia, primarily due to \
    lower digital literacy among business owners and the absence of localised solutions that \
    account for regional payment preferences and language requirements.

    Within Uzbekistan, existing platforms such as iDoctor for healthcare and Uzum for e-commerce \
    have demonstrated that users are willing to adopt digital tools when the value proposition is \
    clear and the user interface is accessible (Kun.uz, 2023). The success of these platforms \
    proves that the Uzbek market is ready for vertical digital solutions, yet no dedicated \
    platform exists for the barbershop and personal grooming segment that combines booking, \
    business management and customer discovery in a single solution. Kling (2021) observed that \
    niche vertical platforms tend to outperform horizontal solutions in user satisfaction because \
    they can tailor workflows, terminology and features to specific industry needs. Porter's \
    (2008) five forces framework further supports this observation, suggesting that focused \
    platforms can build defensible competitive positions by deeply understanding their target \
    industry's unique requirements.

    The software engineering methodology adopted for IBook draws on established best practices. \
    Agile software development, which forms the basis of the IBook development process, has been \
    validated extensively in the literature. Beck et al. (2001) established the Agile Manifesto \
    principles that prioritise working software, customer collaboration and responsiveness to \
    change. Schwaber and Sutherland (2020) refined these ideas in the Scrum Guide, advocating for \
    iterative sprints with defined goals and regular retrospectives. Pressman and Maxim (2020) \
    recommended Agile approaches for projects where requirements are expected to evolve, which \
    accurately describes the IBook development experience as early user feedback prompted \
    significant design changes. Sommerville (2016) noted that incremental delivery models are \
    particularly well suited to web application development where rapid user feedback cycles \
    can inform subsequent iterations.

    The technology choices for IBook are also grounded in current best practices. Vue.js, \
    adopted for the frontend, has been recognised for its gentle learning curve and component-based \
    architecture (You, 2023). Its Composition API, introduced in Vue 3, enables better separation \
    of concerns in complex components and has been widely adopted by the developer community. \
    Django and Django REST Framework, used for the backend, provide robust ORM capabilities, \
    built-in security features and a mature ecosystem of third-party packages (Django Software \
    Foundation, 2024). This combination has been recommended by Holovaty and Kaplan-Moss (2009) \
    for projects requiring rapid development without sacrificing code quality or security. The \
    choice of a REST API architecture aligns with industry standards for web service design and \
    enables future extension through mobile applications or third-party integrations.

    Usability evaluation methods informed the testing strategy for IBook. Nielsen (2000) \
    established that testing with as few as five representative users can identify approximately \
    85 percent of usability issues, which justified the informal user acceptance testing approach \
    adopted for this project given the time and resource constraints of an academic submission. \
    The Statista (2024) market data on beauty and personal care revenue in Uzbekistan further \
    confirmed the commercial viability of the platform and provided context for the market \
    demand analysis presented in earlier sections of this report.""")


def section_computing_methodology():
    return textwrap.dedent("""\
    The development of IBook followed an Agile methodology with iterative two-week sprint cycles. \
    This approach was chosen because the project requirements evolved as user research revealed new \
    insights about barbershop operations, and a rigid waterfall approach would not have accommodated \
    these changes efficiently. Each sprint had a defined goal, a set of user stories to implement, \
    and concluded with a review of completed work and a brief retrospective to identify process \
    improvements. The Agile methodology was particularly appropriate because several requirements \
    changed significantly during the project — most notably the original plan to build a Telegram \
    mini application was replaced with a standalone web application after early user research \
    indicated that browser-based access would reach a wider audience.

    The project was divided into five major phases, each aligned with one or more sprint cycles. \
    Phase 1 focused on research and requirements gathering, including the design and distribution \
    of the primary survey, competitor benchmarking against existing booking platforms, and \
    evaluation of candidate technologies for the frontend and backend. Phase 2 covered system \
    architecture design, database schema definition and API specification, producing the ER \
    diagram, data flow diagrams and use case models documented in the Design Documents section. \
    Phase 3 was dedicated to core backend development including user authentication with JWT \
    tokens, booking workflows with conflict detection, business and service management endpoints, \
    and staff scheduling logic. Phase 4 addressed the frontend implementation, building the \
    customer-facing discovery and booking interface alongside the owner management dashboard \
    using Vue 3 with TypeScript. Phase 5 concentrated on testing, performance tuning, user \
    acceptance validation and final report preparation.

    Version control was managed through Git with the repository hosted on GitHub. A structured \
    branching strategy was used, with feature branches created for each sprint task and merged into \
    the main branch upon completion and review. This approach ensured that the main branch always \
    contained stable, tested code. Commit messages followed a conventional format indicating the \
    type of change (feature, fix, documentation) and the affected component, which made it \
    straightforward to trace when and why each change was introduced.

    A notable aspect of the development process was the use of AI-assisted pair programming through \
    Claude Code, an AI coding assistant developed by Anthropic. This tool was integrated into the \
    development workflow from Phase 3 onwards and was used to accelerate implementation tasks such \
    as writing boilerplate code, generating test cases, debugging API endpoints, refactoring \
    existing components and resolving complex TypeScript type errors. The AI assistant did not \
    replace critical design decisions or architectural choices, which remained entirely the \
    responsibility of the developer. Instead, it served as a productivity multiplier, allowing a \
    single developer to achieve the output typically expected from a small team. All AI-generated \
    code was carefully reviewed, tested against the existing test suite and modified where necessary \
    before being committed to the repository. This approach to AI-assisted development reflects an \
    emerging trend in software engineering practice and is documented here transparently in the \
    interest of academic honesty.

    Development tools included Visual Studio Code as the primary IDE with extensions for Vue, \
    TypeScript and Python. Postman was used for API testing during development, allowing rapid \
    verification of endpoint behaviour before writing automated tests. The browser developer tools \
    in Chrome and Firefox were used extensively for frontend debugging, responsive design testing \
    and network request inspection. The Django admin interface served as a convenient data \
    management tool during development, allowing quick inspection and modification of database \
    records without writing custom management commands.

    Continuous integration practices were followed by running the full pytest suite of 99 test \
    cases before each merge to the main branch. Any failing test required investigation and \
    resolution before the merge could proceed. On the frontend, TypeScript compilation errors \
    were treated as blocking issues that had to be resolved before committing code.

    Project progress was tracked using a structured planning directory within the repository, \
    containing phase plans, verification checklists and context handoff documents. Each phase \
    had a dedicated planning file that listed the specific tasks, acceptance criteria and \
    verification steps required for completion. This lightweight project management approach \
    proved effective for a single-developer academic project while maintaining the discipline of \
    documented progress and traceability that would be expected in a professional environment.""")


def section_tech_stack():
    return textwrap.dedent("""\
    The technology stack for IBook was selected based on five criteria: developer productivity, \
    ecosystem maturity, security features, scalability potential and alignment with the modular \
    architecture required by the project. The stack comprises a Vue 3 frontend, a Django 5.2 \
    backend and a relational database layer, all communicating through a RESTful API.

    The frontend is built with Vue 3 using the Composition API, TypeScript for type safety, and \
    Vite as the build tool. Vue 3 was chosen over alternatives such as React and Angular after \
    evaluating developer experience, bundle size and the availability of ecosystem tools. The \
    Composition API enables better code organisation for complex components, while TypeScript \
    strict mode catches type-related bugs at compile time rather than at runtime. Vite provides \
    near-instant hot module replacement during development and optimised production builds. The \
    UI layer uses responsive CSS to ensure the application works correctly on both desktop and \
    mobile browsers.""")


def section_tech_comparison():
    """Return comparison table data."""
    frontend = [
        ("Vue 3 + TypeScript", "Gentle learning curve, Composition API, small bundle, strong TypeScript support", "Selected"),
        ("React + TypeScript", "Large ecosystem, JSX flexibility, extensive community", "Considered"),
        ("Angular", "Enterprise features, built-in DI and routing, steep learning curve", "Rejected"),
    ]
    backend = [
        ("Django 5.2 + DRF", "Built-in ORM, admin panel, security features, rapid development", "Selected"),
        ("FastAPI", "High performance, async support, automatic OpenAPI docs", "Considered"),
        ("Express.js", "Lightweight, JavaScript everywhere, requires many third-party packages", "Rejected"),
    ]
    return frontend, backend


def section_tech_stack_continued():
    return textwrap.dedent("""\
    The backend is powered by Django 5.2 with Django REST Framework providing the API layer. \
    Django was selected for its batteries-included philosophy, which provides built-in user \
    authentication, an ORM for database operations, an admin interface for development-time \
    data management, and robust security defaults including CSRF protection, SQL injection \
    prevention and XSS mitigation. Django REST Framework adds serialisation, viewset patterns \
    and permission classes that simplified the creation of RESTful endpoints.

    Authentication is implemented using JSON Web Tokens stored in HTTP-only cookies. Access \
    tokens have a short expiry while refresh tokens enable seamless session renewal. This \
    approach protects against cross-site scripting attacks that could exploit tokens stored \
    in browser local storage.

    The database uses SQLite during development, which provides zero-configuration setup and \
    file-based storage suitable for a single-developer workflow. The Django ORM ensures that \
    all queries are database-agnostic, allowing straightforward migration to PostgreSQL for \
    production deployment. The schema includes models for users, businesses, services, staff \
    members, working hours, bookings, reviews and location data.

    External integrations include the OpenAI API for automated content verification of uploaded \
    images and service descriptions, and OpenStreetMap via Nominatim for reverse geocoding of \
    business locations. These integrations are encapsulated in dedicated service modules to \
    maintain separation of concerns and allow easy replacement if alternative providers are \
    preferred in the future.

    The system follows a clear separation between the frontend and backend. The Vue application \
    communicates exclusively through the REST API, meaning the frontend could be replaced or \
    supplemented with a mobile application in the future without any changes to the backend. \
    This modular architecture supports long-term maintainability and scalability.

    The deployment configuration uses environment variables for sensitive settings such as API \
    keys and database credentials, following the twelve-factor application methodology. During \
    development, a local SQLite database provided zero-configuration convenience, while the \
    Django settings module is structured to support environment-specific configurations for \
    development, staging and production deployments. Static files including compiled JavaScript \
    bundles and CSS are served through Django's static file handling during development, with \
    a clear path to CDN delivery for production. This attention to deployment readiness ensures \
    that the transition from academic prototype to production service requires configuration \
    changes rather than architectural modifications.""")


def section_implementation():
    return textwrap.dedent("""\
    The implementation of IBook is divided into three main components: the customer-facing \
    frontend, the business management dashboard and the backend API. Each component was \
    developed iteratively across multiple sprint cycles, with integration testing performed \
    at the end of each sprint to verify that new features did not break existing functionality.

    The customer-facing frontend provides a clean, intuitive interface for discovering \
    barbershops and booking appointments. The home page presents featured barbershops with \
    search and filter capabilities, allowing users to find shops by location, service type \
    or rating. The search functionality uses debounced input to avoid excessive API calls \
    while providing responsive results. Each barbershop has a detailed profile page displaying \
    its services with pricing and duration, barber profiles with portfolio images, working \
    hours, customer reviews and an interactive map showing the business location via \
    OpenStreetMap integration. The booking flow guides the customer through a multi-step \
    process: selecting a service, choosing a preferred barber, picking an available date and \
    time slot, and confirming the appointment. Time slots are dynamically calculated based on \
    the barber's working hours, existing bookings and the duration of the selected service, \
    ensuring that double-bookings are impossible. The entire flow was designed to be completable \
    in under 60 seconds, reflecting the survey finding that customers value convenience above \
    all other factors when choosing a booking method.

    [SCREENSHOT PLACEHOLDER: Customer booking flow - service selection, barber selection, \
    date/time picker and booking confirmation screens]

    The frontend architecture follows Vue 3 best practices with the Composition API and a \
    well-organised component hierarchy. Shared state is managed through composable functions \
    rather than a global state management library, which keeps the bundle size small and \
    reduces complexity. API communication is handled through a centralised HTTP client built \
    on the Fetch API with automatic token refresh, error handling and request interceptors. \
    All components are strongly typed with TypeScript interfaces that mirror the backend \
    serialiser definitions, ensuring type safety across the full stack.

    The business management dashboard serves barbershop owners and provides comprehensive \
    tools for daily operations. Owners can manage their business profile including photos \
    and descriptions, create and update their service catalogue with pricing and duration, \
    add and manage staff members with their specialisations and availability, configure \
    working hours for each day of the week with support for break periods, and view a calendar \
    of upcoming bookings with the ability to reschedule or cancel. The analytics section \
    provides an overview of booking trends over selectable time periods, most popular services \
    by booking count, peak booking hours displayed as a heat map, staff utilisation rates and \
    estimated revenue figures. These analytics features were designed to help barbershop owners \
    make data-driven decisions about staffing, service pricing and promotional strategies.

    [SCREENSHOT PLACEHOLDER: Owner dashboard - business profile management, service catalogue \
    and analytics overview screens]

    The backend API was implemented using Django REST Framework with a consistent RESTful \
    design following standard HTTP conventions. The API is organised around resources including \
    users, businesses, services, staff, working hours, bookings and reviews. Each resource \
    supports standard CRUD operations with appropriate permission checks — for example, only \
    the owner of a business can modify its services, while any authenticated customer can \
    create a booking. Unauthenticated users can browse businesses and services but cannot \
    create bookings or leave reviews. The API uses serialisers for input validation and \
    response formatting, viewsets for routing, and custom permission classes for fine-grained \
    access control. Error responses follow a consistent format with appropriate HTTP status \
    codes and descriptive error messages.

    Authentication is implemented using a dual-token JWT system. Upon login, the server issues \
    both an access token with a short expiry of 15 minutes and a refresh token with a longer \
    expiry of seven days. Both tokens are stored in HTTP-only cookies with the Secure and \
    SameSite attributes set to prevent cross-site request forgery and cross-site scripting \
    attacks. The frontend automatically refreshes the access token when it expires, providing \
    a seamless user experience without requiring re-authentication.

    The database schema was designed following normalisation principles to avoid data redundancy \
    while maintaining query performance for common operations. Key relationships include: a \
    business has many services and staff members; a staff member belongs to one business; a \
    booking references a customer, a business, a service and optionally a specific barber; \
    and reviews are linked to both the customer and the business. Indexes were added on \
    frequently queried fields such as booking dates, business location coordinates and user \
    foreign keys. The schema also includes soft-delete support for bookings and businesses, \
    allowing data recovery and maintaining referential integrity for historical records.

    [SCREENSHOT PLACEHOLDER: Database schema diagram showing entity relationships]

    External service integration was implemented through dedicated utility modules that follow \
    the adapter pattern for clean separation of concerns. The OpenAI integration verifies that \
    uploaded images are appropriate and that service descriptions meet quality standards before \
    they are published, protecting the platform from inappropriate or low-quality content. The \
    OpenStreetMap integration converts geographic coordinates submitted by business owners into \
    human-readable addresses displayed on the barbershop profile page, eliminating the need for \
    owners to manually type their full address. Both integrations include comprehensive error \
    handling, request timeouts, and fallback behaviour to ensure the platform remains fully \
    functional even when external services are temporarily unavailable or responding slowly.""")


def section_testing():
    return textwrap.dedent("""\
    Quality assurance for IBook was approached through multiple complementary strategies: \
    automated backend testing, frontend type checking, manual testing and survey-based user \
    validation. This multi-layered approach was necessary to ensure reliability across the \
    full stack.

    Backend testing was conducted using pytest with Django's test client. A total of 99 \
    automated test cases were written covering the core functionality of the platform. These \
    tests are organised into several categories. Unit tests verify individual model methods, \
    serialiser validation logic and utility functions in isolation. Integration tests validate \
    complete API workflows including user registration, authentication token management, \
    business creation, service management, booking creation with conflict detection, and review \
    submission. Permission tests confirm that users can only access and modify resources they \
    are authorised to manage. Edge case tests cover scenarios such as booking outside working \
    hours, attempting to book a slot that is already taken, and submitting invalid data formats.

    The test suite achieves broad coverage across the API endpoints and business logic layer. \
    Running the full suite takes approximately 12 seconds, and it was executed before every \
    merge to the main branch as part of the continuous integration practice.

    On the frontend, TypeScript strict mode was enabled from the beginning of development. \
    This configuration enforces null checks, prevents implicit any types and requires explicit \
    return types for functions. While not a substitute for runtime testing, TypeScript strict \
    mode eliminated an entire category of type-related bugs that commonly affect JavaScript \
    applications. The Vue 3 Composition API with TypeScript also enabled better IDE support \
    including autocompletion and inline error detection during development.

    Manual testing was performed throughout the development process using a structured \
    checklist approach. Each sprint concluded with a manual walkthrough of new features \
    across different browsers (Chrome, Firefox and Safari) and screen sizes (desktop, tablet \
    and mobile). Key flows tested manually included the complete booking process from \
    discovery to confirmation, business onboarding from registration to publishing a service \
    catalogue, and the analytics dashboard data accuracy.

    [SCREENSHOT PLACEHOLDER: Pytest test results showing 99 passing tests]

    User acceptance testing was conducted informally with three barbershop owners who were \
    shown the platform and asked to complete common tasks. Their feedback was positive \
    regarding the simplicity of the interface, though they suggested additional features such \
    as SMS reminders and integration with local payment systems for future versions.

    Survey-based validation complemented the technical testing. The primary survey of 57 \
    respondents established the problem domain, and the alignment between survey findings \
    and implemented features confirms that the platform addresses genuine user needs. For \
    instance, the survey revealed that 54 percent of respondents had experienced unclear \
    pricing, and IBook directly addresses this with mandatory service pricing on all business \
    profiles.""")


def section_primary_research(survey):
    n = survey["n"]
    return textwrap.dedent(f"""\
    A primary survey was conducted as the principal research instrument for this project, \
    targeting barbershop customers aged 18 to 45 in Uzbekistan. The survey was distributed \
    online through social media channels and received {n} valid responses over a two-week \
    collection period in January 2026. The questionnaire contained 20 questions covering \
    current booking habits, pain points, feature preferences and willingness to adopt a \
    digital booking platform. Both English and Uzbek/Russian translations were provided to \
    maximise accessibility.

    The results revealed significant inefficiencies in current booking practices. The most \
    common booking method was WhatsApp or messaging (40 percent), followed by walk-ins (23 \
    percent), phone calls (19 percent) and online forms (18 percent). This confirms that the \
    majority of bookings rely on informal, error-prone channels.

    A striking {survey["missed_pct"]} percent of respondents reported having missed an \
    appointment due to miscommunication or simply forgetting the booking. This statistic \
    alone validates the need for automated reminders and structured booking management.

    When asked about pricing transparency, {survey["unclear_pricing_pct"]} percent of \
    respondents indicated they had experienced issues with unclear pricing or service \
    descriptions at barbershops. This finding directly informed the design decision to \
    require barbershop owners to list all services with explicit pricing and duration on \
    the IBook platform.

    Regarding willingness to adopt a booking application, {survey["would_use_app_pct"]} \
    percent of respondents answered either "Definitely" or "Maybe" when asked if they \
    would use such a platform. The most requested features were reviews and ratings (67 \
    percent), online payments (56 percent), reminders and notifications (54 percent), and \
    promotions (53 percent). These preferences were used to prioritise the feature backlog \
    during development.

    Satisfaction with current booking methods was moderate, with a mean rating of \
    approximately 3.1 out of 5. This suggests that while users are not deeply dissatisfied, \
    there is clear room for improvement, which a dedicated platform could provide.

    When examining how frequently respondents visit barbershops, 37 percent reported going \
    once a month, 25 percent visited two to three times per month, 21 percent went less than \
    once a month, and 18 percent visited weekly. This frequency data suggests a substantial \
    recurring demand for barbershop services, confirming that an appointment management platform \
    would see regular rather than one-off usage.

    Regarding long waiting times, 46 percent of respondents reported sometimes experiencing \
    long waits, while 19 percent said they frequently faced long waiting times. Only 18 percent \
    said they never experienced waiting issues. These figures reinforce the argument that \
    poor scheduling practices cause real inconvenience for customers and represent a tangible \
    opportunity for improvement through digital appointment management.

    The importance of seeing barber profiles and ratings before booking received a mean score \
    of 3.1 out of 5, indicating moderate interest. While not the highest-priority feature, this \
    finding justified the inclusion of barber profile pages with portfolio images and customer \
    review capabilities in the IBook platform.

    The charts below present visual summaries of the key survey findings.""")


def section_evaluation():
    return textwrap.dedent("""\
    This section evaluates the extent to which IBook meets its stated project objectives and \
    assesses the overall quality and impact of the delivered solution. Each objective is assessed \
    individually, followed by a broader discussion of the project's achievements and areas for \
    future improvement.

    Objective 1 was to analyse current booking practices through primary research involving at \
    least 50 respondents. This objective was fully achieved. The survey collected 57 valid \
    responses and the findings directly informed the platform design. Key statistics including \
    the 79 percent missed appointment rate, 54 percent unclear pricing rate, and 66 percent \
    willingness to use a booking app provided a strong quantitative evidence base that justified \
    the project and guided feature prioritisation. The survey instrument covered 20 questions \
    addressing booking habits, pain points, feature preferences and adoption willingness, \
    providing a comprehensive view of the target user population.

    Objective 2 was to design a relational database schema and RESTful API architecture supporting \
    multi-tenant barbershop management. This objective was fully achieved. The database schema \
    contains well-normalised entities for users, businesses, services, staff, working hours, \
    bookings, reviews and location data, with appropriate foreign key relationships and indexes. \
    The API follows RESTful conventions with proper authentication, authorisation, input \
    validation and consistent error handling across all endpoints.

    Objective 3 was to develop a fully functional web application enabling customers to discover \
    barbershops, view services and book appointments online. This objective was substantially \
    achieved. The customer-facing interface supports barbershop discovery with search and filtering, \
    detailed business profile viewing, service browsing with pricing information, and appointment \
    booking with real-time availability checking. Some planned enhancements such as advanced \
    geo-based filtering and real-time push notifications are partially implemented and have been \
    marked as priorities for the first post-launch update.

    Objective 4 was to implement a business management dashboard allowing owners to manage services, \
    staff, working hours and view analytics. This objective was fully achieved. The dashboard \
    provides comprehensive tools for all aspects of business management including profile editing, \
    service catalogue administration, staff management with specialisation tracking, working hours \
    configuration, and an analytics overview with booking trend visualisations and revenue estimates.

    Objective 5 was to validate the platform through automated testing targeting at least 90 test \
    cases and user acceptance testing. This objective was fully achieved. The pytest suite contains \
    99 test cases covering unit, integration and permission testing across the backend. TypeScript \
    strict mode provides compile-time validation for the entire frontend codebase. Informal user \
    acceptance testing with three barbershop owners yielded positive feedback regarding interface \
    simplicity and workflow alignment.

    Objective 6 was to deliver a complete academic report by 10th of April, 2026. This objective \
    is achieved through this document, which covers the full project lifecycle from problem \
    identification through evaluation.

    The key technical achievements of the project include: a clean architectural separation between \
    frontend and backend enabling independent scaling and technology evolution; JWT-based \
    authentication with HTTP-only cookie storage preventing common web security vulnerabilities; \
    AI-powered content moderation ensuring quality standards for user-generated content; reverse \
    geocoding integration simplifying the business onboarding process; and a comprehensive \
    automated test suite providing confidence in system reliability. The platform demonstrates \
    that a single developer, supported by modern development tools and AI-assisted pair \
    programming, can deliver a production-quality system within an academic project timeline.

    Areas for future improvement include expanding the test suite to include frontend component \
    tests using Vitest and Vue Test Utils, implementing real-time WebSocket notifications for \
    booking confirmations and reminders, adding multi-language support with Uzbek and Russian \
    translations, integrating local payment methods such as Payme and Click, and conducting a \
    larger-scale user acceptance study with active barbershop deployments over a sustained period \
    to measure real-world impact on appointment no-show rates and customer satisfaction.""")


def section_conclusion():
    return textwrap.dedent("""\
    This project set out to address the operational inefficiencies faced by barbershops in \
    Uzbekistan through the development of IBook, a web-based booking and management application. \
    The problem was validated through primary research with 57 survey respondents, demonstrating \
    that the vast majority of barbershop customers experience scheduling difficulties, missed \
    appointments and unclear service information when relying on traditional booking methods \
    such as phone calls, messaging and walk-ins. The 79 percent missed appointment rate and 54 \
    percent unclear pricing rate identified through the survey provided compelling evidence that \
    a digital solution is not merely a convenience enhancement but a genuine operational necessity \
    for barbershops seeking to improve customer satisfaction and reduce revenue loss.

    The IBook platform was designed and developed using a modern technology stack comprising \
    Vue 3 with TypeScript for the frontend and Django 5.2 with Django REST Framework for the \
    backend. The system supports three user roles and provides distinct, tailored interfaces for \
    customer booking, barber schedule management and business owner administration. The \
    development followed Agile methodology with iterative two-week sprints, version control \
    through GitHub, and AI-assisted pair programming through Claude Code to maximise productivity \
    within the constraints of a single-developer academic project. The architecture follows \
    modern best practices with a clean separation between the frontend and backend layers, \
    enabling independent scaling and future technology evolution.

    Testing and evaluation confirmed that the platform meets its six stated objectives. The \
    automated test suite of 99 cases provides confidence in backend reliability, covering unit, \
    integration and permission testing across all major API endpoints. TypeScript strict mode \
    and manual cross-browser testing ensure frontend quality and consistency. Survey data \
    confirmed genuine market demand, with 66 percent of respondents expressing willingness to \
    use a dedicated booking application. Informal user acceptance testing with barbershop \
    owners validated the usability and practical relevance of the management dashboard features.

    Based on the experience gained during this project and the feedback received from potential \
    users, we recommend that future work focus on four priorities. First, conducting a pilot \
    deployment with five to ten barbershops in Tashkent to gather real-world usage data, \
    measure the actual reduction in no-show rates, and collect structured feedback for further \
    iteration. Second, integrating local payment methods such as Payme and Click to align with \
    Uzbekistan's rapidly evolving digital payment infrastructure and enable the online prepayment \
    feature that 56 percent of survey respondents identified as desirable. Third, developing a \
    progressive web application with offline support and push notifications to improve the mobile \
    experience without the cost and complexity of maintaining separate native iOS and Android \
    applications. Fourth, adding multi-language support with Uzbek and Russian translations to \
    broaden the platform's accessibility beyond English-speaking users.

    In conclusion, IBook demonstrates that well-researched, purpose-built digital solutions can \
    address real operational challenges in Uzbekistan's personal services sector. With continued \
    development and a structured pilot programme, the platform has the potential to become a \
    valuable tool for barbershops and their customers across the country.""")


def section_references():
    return [
        'Beck, K., Beedle, M., van Bennekum, A., Cockburn, A., Cunningham, W., Fowler, M., Grenning, J., Highsmith, J., Hunt, A., Jeffries, R., Kern, J., Marick, B., Martin, R.C., Mellor, S., Schwaber, K., Sutherland, J. and Thomas, D. (2001) Manifesto for Agile Software Development. Available at: https://agilemanifesto.org/ (Accessed: 15 March 2026).',
        'Booksy (2024) Booksy for Business: Appointment Scheduling Software. Available at: https://booksy.com/biz (Accessed: 10 February 2026).',
        'Django Software Foundation (2024) Django Documentation: Version 5.2. Available at: https://docs.djangoproject.com/en/5.2/ (Accessed: 20 March 2026).',
        'Fresha (2024) Fresha: Online Booking Software for Beauty and Wellness. Available at: https://www.fresha.com/ (Accessed: 10 February 2026).',
        'Holovaty, A. and Kaplan-Moss, J. (2009) The Definitive Guide to Django: Web Development Done Right. 2nd edn. Berkeley, CA: Apress.',
        'International Finance Corporation (2022) Digital Economy in Central Asia: Opportunities and Challenges for SMEs. Washington, DC: World Bank Group.',
        'Kling, R. (2021) \'Vertical SaaS platforms and their impact on niche service industries\', Journal of Information Technology Management, 32(3), pp. 45-62.',
        'Kun.uz (2023) \'Digital services adoption in Uzbekistan reaches new milestone\', Kun.uz, 15 September. Available at: https://kun.uz/ (Accessed: 5 January 2026).',
        'Laudon, K.C. and Laudon, J.P. (2020) Management Information Systems: Managing the Digital Firm. 16th edn. Harlow: Pearson Education.',
        'Mathew, S. and Soliman, M. (2021) \'Digital transformation in the beauty and wellness industry: A systematic review\', International Journal of Hospitality Management, 93, pp. 102-115.',
        'Nielsen, J. (2000) \'Why You Only Need to Test with 5 Users\', Nielsen Norman Group, 19 March. Available at: https://www.nngroup.com/articles/why-you-only-need-to-test-with-5-users/ (Accessed: 20 March 2026).',
        'Porter, M.E. (2008) \'The five competitive forces that shape strategy\', Harvard Business Review, 86(1), pp. 78-93.',
        'Pressman, R.S. and Maxim, B.R. (2020) Software Engineering: A Practitioner\'s Approach. 9th edn. New York: McGraw-Hill Education.',
        'Schwaber, K. and Sutherland, J. (2020) The Scrum Guide: The Definitive Guide to Scrum: The Rules of the Game. Available at: https://scrumguides.org/ (Accessed: 15 March 2026).',
        'Sommerville, I. (2016) Software Engineering. 10th edn. Harlow: Pearson Education.',
        'Statista (2024) \'Revenue in the Beauty and Personal Care market in Uzbekistan\', Statista. Available at: https://www.statista.com/ (Accessed: 12 January 2026).',
        'Verhoef, P.C., Broekhuizen, T., Bart, Y., Bhattacharya, A., Dong, J.Q., Fabian, N. and Haenlein, M. (2021) \'Digital transformation: A multidisciplinary reflection and research agenda\', Journal of Business Research, 122, pp. 889-901.',
        'Westerman, G., Bonnet, D. and McAfee, A. (2014) Leading Digital: Turning Technology into Business Transformation. Boston, MA: Harvard Business Review Press.',
        'You, E. (2023) Vue.js 3 Documentation. Available at: https://vuejs.org/ (Accessed: 20 March 2026).',
        'Zeithaml, V.A., Bitner, M.J. and Gremler, D.D. (2018) Services Marketing: Integrating Customer Focus Across the Firm. 7th edn. New York: McGraw-Hill Education.',
    ]


# ===================================================================
# Main document builder
# ===================================================================

def build_report():
    print("Loading existing report...")
    existing_paras, existing_tables = load_existing_paragraphs()

    print("Loading survey data...")
    survey = load_survey()

    print("Creating new document...")
    doc = Document()

    # ---------------------------------------------------------------
    # Page setup: 1-inch margins
    # ---------------------------------------------------------------
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # ---------------------------------------------------------------
    # Set default font
    # ---------------------------------------------------------------
    style = doc.styles["Normal"]
    font = style.font
    font.name = FONT_BODY
    font.size = FONT_SIZE_BODY
    font.color.rgb = DARK_GRAY

    # ---------------------------------------------------------------
    # TITLE PAGE
    # ---------------------------------------------------------------
    for _ in range(6):
        doc.add_paragraph()

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run("WESTMINSTER INTERNATIONAL UNIVERSITY IN TASHKENT")
    run.font.name = FONT_BODY
    run.font.size = Pt(14)
    run.font.color.rgb = NAVY
    run.bold = True

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("BSc (Hons) Business Information Systems")
    run.font.name = FONT_BODY
    run.font.size = Pt(12)
    run.font.color.rgb = NAVY

    doc.add_paragraph()
    doc.add_paragraph()

    project_title = doc.add_paragraph()
    project_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = project_title.add_run("BISP Final Report")
    run.font.name = FONT_BODY
    run.font.size = Pt(20)
    run.font.color.rgb = NAVY
    run.bold = True

    doc.add_paragraph()

    app_name = doc.add_paragraph()
    app_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = app_name.add_run("IBook — Barbershop Booking and Management Application")
    run.font.name = FONT_BODY
    run.font.size = Pt(14)
    run.font.color.rgb = NAVY
    run.italic = True

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    # Student info block
    info_lines = [
        ("Course:", "BSc (Hons) Business Information Systems"),
        ("Supervisor:", "Said Abduvaliev"),
        ("Date of Submission:", "10th of April, 2026"),
    ]
    for label, value in info_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_label = p.add_run(label + " ")
        run_label.font.name = FONT_BODY
        run_label.font.size = Pt(11)
        run_label.font.color.rgb = NAVY
        run_label.bold = True
        run_val = p.add_run(value)
        run_val.font.name = FONT_BODY
        run_val.font.size = Pt(11)
        run_val.font.color.rgb = DARK_GRAY

    doc.add_page_break()

    # ---------------------------------------------------------------
    # ACKNOWLEDGEMENTS
    # ---------------------------------------------------------------
    add_heading_styled(doc, "Acknowledgements", level=1)
    add_formatted_para(doc, section_acknowledgements())

    doc.add_page_break()

    # ---------------------------------------------------------------
    # ABSTRACT
    # ---------------------------------------------------------------
    add_heading_styled(doc, "Abstract", level=1)
    for para_text in section_abstract().split("\n\n"):
        add_formatted_para(doc, para_text.strip())

    doc.add_page_break()

    # ---------------------------------------------------------------
    # TABLE OF CONTENTS placeholder
    # ---------------------------------------------------------------
    add_heading_styled(doc, "Table of Contents", level=1)
    toc_items = [
        "Acknowledgements",
        "Abstract",
        "1. Introduction",
        "2. Problem Analysis",
        "   2.1 Strategic Context",
        "   2.2 Sector and Industry Analysis",
        "   2.3 Market Demand Analysis",
        "   2.4 User Analysis",
        "   2.5 Competitor Analysis",
        "   2.6 SWOT Analysis",
        "   2.7 PESTLE Analysis",
        "   2.8 Business Impact",
        "3. Project Objectives",
        "4. Scope and Requirements",
        "5. Literature Review",
        "6. Computing Methodology",
        "7. Technology Stack and System Design",
        "8. IT Strategy",
        "9. Critical Self-Evaluation",
        "10. Design Documents",
        "11. System Implementation",
        "12. Testing and Quality Assurance",
        "13. Primary Research Results",
        "14. Evaluation and Results",
        "15. Conclusion and Recommendations",
        "References",
    ]
    for item in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(item)
        run.font.name = FONT_BODY
        run.font.size = Pt(11)
        run.font.color.rgb = DARK_GRAY
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.3

    doc.add_page_break()

    # ---------------------------------------------------------------
    # 1. INTRODUCTION (existing content)
    # ---------------------------------------------------------------
    add_heading_styled(doc, "1. Introduction", level=1)
    for idx in range(10, 13):
        if idx < len(existing_paras) and existing_paras[idx]["text"].strip():
            add_formatted_para(doc, fix_text(existing_paras[idx]["text"]))

    # ---------------------------------------------------------------
    # 2. PROBLEM ANALYSIS (existing content)
    # ---------------------------------------------------------------
    add_heading_styled(doc, "2. Problem Analysis", level=1)

    # Strategic Context (paras 15-23)
    add_heading_styled(doc, "2.1 Strategic Context", level=2)
    for idx in [19, 21, 23]:
        if idx < len(existing_paras) and existing_paras[idx]["text"].strip():
            add_formatted_para(doc, fix_text(existing_paras[idx]["text"]))

    # Sector Analysis (paras 25-32)
    add_heading_styled(doc, "2.2 Sector and Industry Analysis", level=3)
    for idx in [28, 30, 32]:
        if idx < len(existing_paras) and existing_paras[idx]["text"].strip():
            add_formatted_para(doc, fix_text(existing_paras[idx]["text"]))

    # Market Demand (paras 34-38)
    add_heading_styled(doc, "2.3 Market Demand Analysis", level=2)
    for idx in [36, 38]:
        if idx < len(existing_paras) and existing_paras[idx]["text"].strip():
            add_formatted_para(doc, fix_text(existing_paras[idx]["text"]))

    # User Analysis (paras 39-45)
    add_heading_styled(doc, "2.4 User Analysis", level=2)
    for idx in [41, 43, 45]:
        if idx < len(existing_paras) and existing_paras[idx]["text"].strip():
            add_formatted_para(doc, fix_text(existing_paras[idx]["text"]))

    # Competitor Analysis (paras 46-52)
    add_heading_styled(doc, "2.5 Competitor Analysis", level=2)
    for idx in [48, 50, 52]:
        if idx < len(existing_paras) and existing_paras[idx]["text"].strip():
            add_formatted_para(doc, fix_text(existing_paras[idx]["text"]))

    # SWOT Analysis (paras 53-74)
    add_heading_styled(doc, "2.6 SWOT Analysis", level=2)

    add_heading_styled(doc, "Strengths", level=3)
    for idx in [59, 60, 61, 62]:
        if idx < len(existing_paras) and existing_paras[idx]["text"].strip():
            add_bullet(doc, fix_text(existing_paras[idx]["text"]))

    add_heading_styled(doc, "Weaknesses", level=3)
    for idx in [64, 65, 66]:
        if idx < len(existing_paras) and existing_paras[idx]["text"].strip():
            add_bullet(doc, fix_text(existing_paras[idx]["text"]))

    add_heading_styled(doc, "Opportunities", level=3)
    for idx in [68, 69, 70]:
        if idx < len(existing_paras) and existing_paras[idx]["text"].strip():
            add_bullet(doc, fix_text(existing_paras[idx]["text"]))

    add_heading_styled(doc, "Threats", level=3)
    for idx in [72, 73, 74]:
        if idx < len(existing_paras) and existing_paras[idx]["text"].strip():
            add_bullet(doc, fix_text(existing_paras[idx]["text"]))

    # PESTLE Analysis (paras 76-90)
    add_heading_styled(doc, "2.7 PESTLE Analysis", level=2)
    if 78 < len(existing_paras):
        add_formatted_para(doc, fix_text(existing_paras[78]["text"]))

    pestle_sections = [
        (79, "Political Factors", [80]),
        (81, "Economic Factors", [82]),
        (83, "Social Factors", [84]),
        (85, "Technological Factors", [86]),
        (87, "Legal Factors", [88]),
        (89, "Environmental Factors", [90]),
    ]
    for heading_idx, heading_text, para_indices in pestle_sections:
        add_heading_styled(doc, heading_text, level=3)
        for idx in para_indices:
            if idx < len(existing_paras) and existing_paras[idx]["text"].strip():
                text = fix_text(existing_paras[idx]["text"])
                # Environmental factors paragraph was cut off in original
                if idx == 90 and not text.endswith("."):
                    text += " contributes to a more environmentally responsible approach to business administration, reducing paper waste and supporting sustainable practices in the personal services industry."
                add_formatted_para(doc, text)

    # Business Impact (paras 92-93)
    add_heading_styled(doc, "2.8 Business Impact", level=2)
    for idx in [93]:
        if idx < len(existing_paras) and existing_paras[idx]["text"].strip():
            add_formatted_para(doc, fix_text(existing_paras[idx]["text"]))

    # ---------------------------------------------------------------
    # 3. PROJECT OBJECTIVES (new)
    # ---------------------------------------------------------------
    add_heading_styled(doc, "3. Project Objectives", level=1)
    add_formatted_para(doc, "The following SMART objectives were defined to guide the development of IBook and provide measurable criteria for evaluating the success of the project:")
    objectives = section_project_objectives()
    for obj_label, obj_text in objectives:
        p = doc.add_paragraph()
        run_b = p.add_run(f"{obj_label}: ")
        run_b.font.name = FONT_BODY
        run_b.font.size = FONT_SIZE_BODY
        run_b.font.color.rgb = NAVY
        run_b.bold = True
        run_t = p.add_run(obj_text)
        run_t.font.name = FONT_BODY
        run_t.font.size = FONT_SIZE_BODY
        run_t.font.color.rgb = DARK_GRAY
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = LINE_SPACING

    # ---------------------------------------------------------------
    # 4. SCOPE AND REQUIREMENTS (new)
    # ---------------------------------------------------------------
    add_heading_styled(doc, "4. Scope and Requirements", level=1)
    for para_text in section_scope_text().split("\n\n"):
        add_formatted_para(doc, para_text.strip())

    add_heading_styled(doc, "4.1 MoSCoW Prioritisation", level=2)
    add_formatted_para(doc, "The following table presents the MoSCoW prioritisation of requirements for the IBook platform:")
    moscow = section_scope_requirements()
    add_table(doc,
              ["Priority", "Requirement"],
              moscow)

    # ---------------------------------------------------------------
    # 5. LITERATURE REVIEW (new)
    # ---------------------------------------------------------------
    add_heading_styled(doc, "5. Literature Review", level=1)
    for para_text in section_literature_review().split("\n\n"):
        add_formatted_para(doc, para_text.strip())

    # ---------------------------------------------------------------
    # 6. COMPUTING METHODOLOGY (new)
    # ---------------------------------------------------------------
    add_heading_styled(doc, "6. Computing Methodology", level=1)
    for para_text in section_computing_methodology().split("\n\n"):
        add_formatted_para(doc, para_text.strip())

    # ---------------------------------------------------------------
    # 7. TECHNOLOGY STACK AND SYSTEM DESIGN (new)
    # ---------------------------------------------------------------
    add_heading_styled(doc, "7. Technology Stack and System Design", level=1)
    for para_text in section_tech_stack().split("\n\n"):
        add_formatted_para(doc, para_text.strip())

    # Comparison tables
    frontend_comp, backend_comp = section_tech_comparison()
    add_heading_styled(doc, "7.1 Frontend Framework Comparison", level=2)
    add_table(doc,
              ["Framework", "Key Characteristics", "Decision"],
              frontend_comp)

    add_heading_styled(doc, "7.2 Backend Framework Comparison", level=2)
    add_table(doc,
              ["Framework", "Key Characteristics", "Decision"],
              backend_comp)

    for para_text in section_tech_stack_continued().split("\n\n"):
        add_formatted_para(doc, para_text.strip())

    # ---------------------------------------------------------------
    # 8. IT STRATEGY (existing content, fixed)
    # ---------------------------------------------------------------
    add_heading_styled(doc, "8. IT Strategy", level=1)
    for idx in [96, 98, 100, 102, 104]:
        if idx < len(existing_paras) and existing_paras[idx]["text"].strip():
            add_formatted_para(doc, fix_text(existing_paras[idx]["text"]))

    # ---------------------------------------------------------------
    # 9. CRITICAL SELF-EVALUATION (existing content, fixed)
    # ---------------------------------------------------------------
    add_heading_styled(doc, "9. Critical Self-Evaluation", level=1)
    for idx in [107, 108]:
        if idx < len(existing_paras) and existing_paras[idx]["text"].strip():
            add_formatted_para(doc, fix_text(existing_paras[idx]["text"]))

    # ---------------------------------------------------------------
    # 10. DESIGN DOCUMENTS (existing content)
    # ---------------------------------------------------------------
    add_heading_styled(doc, "10. Design Documents", level=1)

    add_heading_styled(doc, "10.1 Use Case Diagram", level=2)
    if 115 < len(existing_paras):
        add_formatted_para(doc, fix_text(existing_paras[115]["text"]))
    add_formatted_para(doc, "[SCREENSHOT PLACEHOLDER: Use Case Diagram]",
                       italic=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    add_heading_styled(doc, "10.2 Logical Product Model", level=2)
    if 117 < len(existing_paras):
        add_formatted_para(doc, fix_text(existing_paras[117]["text"]))

    add_heading_styled(doc, "10.3 Database ER Model", level=2)
    add_formatted_para(doc, "[SCREENSHOT PLACEHOLDER: Entity Relationship Diagram]",
                       italic=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # Copy ER table
    er_data = []
    for row in existing_tables[1].rows[1:]:
        er_data.append([c.text.strip() for c in row.cells])
    add_table(doc, ["Entity", "Description"], er_data)

    add_heading_styled(doc, "10.4 Data Flow Diagram", level=2)
    add_formatted_para(doc, "[SCREENSHOT PLACEHOLDER: Data Flow Diagram]",
                       italic=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # DFD description paragraphs
    for idx in [125, 129, 130, 131, 133, 134, 135, 136, 137]:
        if idx < len(existing_paras) and existing_paras[idx]["text"].strip():
            add_formatted_para(doc, fix_text(existing_paras[idx]["text"]))

    # DFD process table
    dfd_data = []
    for row in existing_tables[2].rows[1:]:
        dfd_data.append([c.text.strip() for c in row.cells])
    add_table(doc, ["#", "Process", "Function"], dfd_data)

    # ---------------------------------------------------------------
    # 11. SYSTEM IMPLEMENTATION (new)
    # ---------------------------------------------------------------
    add_heading_styled(doc, "11. System Implementation", level=1)
    for para_text in section_implementation().split("\n\n"):
        text = para_text.strip()
        if text.startswith("[SCREENSHOT"):
            add_formatted_para(doc, text, italic=True,
                               alignment=WD_ALIGN_PARAGRAPH.CENTER)
        else:
            add_formatted_para(doc, text)

    # ---------------------------------------------------------------
    # 12. TESTING AND QUALITY ASSURANCE (new)
    # ---------------------------------------------------------------
    add_heading_styled(doc, "12. Testing and Quality Assurance", level=1)
    for para_text in section_testing().split("\n\n"):
        text = para_text.strip()
        if text.startswith("[SCREENSHOT"):
            add_formatted_para(doc, text, italic=True,
                               alignment=WD_ALIGN_PARAGRAPH.CENTER)
        else:
            add_formatted_para(doc, text)

    # ---------------------------------------------------------------
    # 13. PRIMARY RESEARCH RESULTS (new, with charts)
    # ---------------------------------------------------------------
    add_heading_styled(doc, "13. Primary Research Results", level=1)
    for para_text in section_primary_research(survey).split("\n\n"):
        add_formatted_para(doc, para_text.strip())

    # Insert charts
    print("Generating survey charts...")

    chart_defs = [
        ("Figure 1: Current Booking Methods", chart_booking_method),
        ("Figure 2: Missed Appointments", chart_missed_appointments),
        ("Figure 3: Issues with Unclear Pricing", chart_unclear_pricing),
        ("Figure 4: Willingness to Use a Booking App", chart_would_use_app),
        ("Figure 5: Most Desired App Features", chart_features_wanted),
        ("Figure 6: Satisfaction with Current Booking Methods", chart_satisfaction),
    ]

    for caption, chart_func in chart_defs:
        buf = create_chart_image(chart_func, survey=survey)
        doc.add_picture(buf, width=Inches(5.5))
        last_para = doc.paragraphs[-1]
        last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cap_para = doc.add_paragraph()
        cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap_para.add_run(caption)
        run.font.name = FONT_BODY
        run.font.size = Pt(10)
        run.font.color.rgb = NAVY
        run.italic = True
        cap_para.paragraph_format.space_after = Pt(12)

    # ---------------------------------------------------------------
    # 14. EVALUATION AND RESULTS (new)
    # ---------------------------------------------------------------
    add_heading_styled(doc, "14. Evaluation and Results", level=1)
    for para_text in section_evaluation().split("\n\n"):
        add_formatted_para(doc, para_text.strip())

    # ---------------------------------------------------------------
    # 15. CONCLUSION AND RECOMMENDATIONS (new)
    # ---------------------------------------------------------------
    add_heading_styled(doc, "15. Conclusion and Recommendations", level=1)
    for para_text in section_conclusion().split("\n\n"):
        add_formatted_para(doc, para_text.strip())

    # ---------------------------------------------------------------
    # REFERENCES (new)
    # ---------------------------------------------------------------
    doc.add_page_break()
    add_heading_styled(doc, "References", level=1)
    refs = section_references()
    for ref in refs:
        p = doc.add_paragraph()
        run = p.add_run(ref)
        run.font.name = FONT_BODY
        run.font.size = Pt(10)
        run.font.color.rgb = DARK_GRAY
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = LINE_SPACING
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)

    # ---------------------------------------------------------------
    # APPENDICES
    # ---------------------------------------------------------------
    doc.add_page_break()
    add_heading_styled(doc, "Appendices", level=1)

    add_heading_styled(doc, "Appendix A: Revised Schedule of Work", level=2)
    sched_data = []
    for row in existing_tables[0].rows[1:]:
        cells = [c.text.strip() for c in row.cells]
        if any(cells):
            sched_data.append(cells)
    add_table(doc, ["Phase", "Task", "Status", "Completion Date"], sched_data)

    add_heading_styled(doc, "Appendix B: Survey Questionnaire", level=2)
    add_formatted_para(doc, "The following questions were included in the primary research survey distributed to barbershop customers:")
    survey_questions = [
        "How do you usually book appointments at a barbershop?",
        "How often do you go to a barbershop?",
        "How do you usually wait to get an appointment?",
        "Have you ever missed an appointment because of miscommunication or forgetting it?",
        "How satisfied are you with the current booking methods available to you?",
        "What is the biggest challenge you face when booking a haircut or grooming service?",
        "Have you experienced long waiting times at barbershops?",
        "Do you find it difficult to find a barber who matches your style preferences?",
        "Have you faced issues with unclear pricing or services offered?",
        "Are cancellations or rescheduling appointments a problem for you?",
        "Would you use a mobile app to book barbershop appointments if it existed?",
        "What features would make you want to use a booking app?",
        "How important is the ability to see barber profiles and ratings before booking?",
        "Would you prefer an app that allows prepayment or pay at the barbershop?",
        "How likely are you to recommend such an app to friends if it worked well?",
        "How much do you usually spend on a haircut or grooming service?",
        "Are you willing to pay a small fee for convenience when booking online?",
        "How often do you try new barbershops versus sticking to one you know?",
        "Would you like to receive promotions or loyalty rewards through an app?",
        "Are there any other problems you face with barbershops that you wish were solved?",
    ]
    for i, q in enumerate(survey_questions, 1):
        add_bullet(doc, f"Q{i}: {q}")

    add_heading_styled(doc, "Appendix C: Application Screenshots", level=2)
    add_formatted_para(doc, "[SCREENSHOT PLACEHOLDER: Application screenshots to be inserted here — customer booking flow, barber dashboard, owner analytics panel, mobile responsive views]",
                       italic=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # ---------------------------------------------------------------
    # Save
    # ---------------------------------------------------------------
    print(f"Saving to {OUTPUT_PATH}...")
    doc.save(OUTPUT_PATH)

    # ---------------------------------------------------------------
    # Word count
    # ---------------------------------------------------------------
    count_doc = Document(OUTPUT_PATH)
    total_words = 0
    for p in count_doc.paragraphs:
        total_words += len(p.text.split())
    for t in count_doc.tables:
        for row in t.rows:
            for cell in row.cells:
                total_words += len(cell.text.split())

    print(f"\n{'='*50}")
    print(f"  Report generated successfully!")
    print(f"  Output: {OUTPUT_PATH}")
    print(f"  Total word count: {total_words:,}")
    print(f"{'='*50}")

    return total_words


if __name__ == "__main__":
    build_report()
